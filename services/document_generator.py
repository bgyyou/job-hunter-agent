"""v3 M-rebuild-3: 文档生成统一接口（Word + PDF）

按 update_plan.md §1.4 / §2.8 / §3.1 硬约束：
- Word：python-docx + jinja2 模板（2 套风格：保守/现代），HR 主流格式可编辑
- PDF：复用 tools/generator/resume_pdf.py（playwright headless chromium），
      保留现状避免引入 weasyprint 等系统库
- 一页纸强校验：调 OnePageEstimator，overflow → 直接抛 ValueError
- 文件命名：`{姓名}_{岗位}_{公司}.{ext}`（特殊字符过滤）

**严禁修改 §1.4 字号/边距/行距硬约束**——这是 v3 一页纸铁律的根基。
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional, Literal

from loguru import logger

# 内部复用 round-1 + v2.1 既有模块
from services.one_page_estimator import OnePageEstimator, PageEstimate
from tools.generator.resume_pdf import html_to_pdf, html_to_pdf_safe


TEMPLATE_DIR = Path(__file__).parent / "document_generator_templates" / "word"
TemplateName = Literal["conservative", "modern"]


# ============================================================
# 自定义异常（捕获更友好）
# ============================================================

class DocumentGenerationError(Exception):
    """文档生成失败（基类）。"""


class OnePageOverflowError(DocumentGenerationError):
    """简历超出一页纸——按 update_plan.md §1.4 不允许超页导出。"""

    def __init__(self, estimate: PageEstimate):
        self.estimate = estimate
        msg = (
            f"简历超出一页纸（{estimate.total_mm:.1f}mm > {estimate.capacity_mm}mm）。\n"
            f"超页段：{', '.join(estimate.overflow_segments) or '无明确段'}\n"
            f"建议：{chr(10).join(estimate.suggestions[:3]) or '精简整体内容'}"
        )
        super().__init__(msg)


# ============================================================
# 工具函数
# ============================================================

def sanitize_filename_part(s: str, fallback: str = "未命名") -> str:
    """过滤文件名字符：保留中文/英文/数字/下划线/连字符，其他替换为 '_'。"""
    if not s or not s.strip():
        return fallback
    # 移除 Windows 非法字符 <>:"/\\|?* 和控制字符
    s = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", s)
    # 多余下划线/空白压一个
    s = re.sub(r"[_\s]+", "_", s).strip("._")
    return s[:40] or fallback


def suggest_filename(name: str, jd_title: str, company: str, ext: str = "docx") -> str:
    """生成导出文件名：`{姓名}_{岗位}_{公司}.{ext}`。

    Args:
        name: 简历姓名
        jd_title: 目标岗位（来自 JD）
        company: 目标公司（来自 JD）
        ext: 文件扩展名（docx / pdf）

    Returns:
        过滤特殊字符后的安全文件名
    """
    parts = [
        sanitize_filename_part(name, "简历"),
        sanitize_filename_part(jd_title, "岗位"),
        sanitize_filename_part(company, "公司"),
    ]
    return f"{'_'.join(parts)}.{ext.lstrip('.')}"


# ============================================================
# DocumentGenerator 主类
# ============================================================

@dataclass
class DocumentResult:
    """生成结果（含文件名 + bytes + 估算信息）。"""

    filename: str
    content: bytes
    mime_type: str
    template_used: str
    estimate: PageEstimate


class DocumentGenerator:
    """统一文档生成器：Word + PDF，内部走 python-docx / jinja2 / playwright。"""

    # §1.4 硬约束（10.5pt + 1.2 行距 + A4 265mm）
    DEFAULT_FONT_SIZE_PT = 10.5
    SECTION_TITLE_SIZE_PT = 12.0
    LINE_SPACING = 1.2
    A4_MARGIN_TOP_MM = 12
    A4_MARGIN_BOTTOM_MM = 12
    A4_MARGIN_LEFT_MM = 14
    A4_MARGIN_RIGHT_MM = 14

    def __init__(self, estimator: Optional[OnePageEstimator] = None):
        self.estimator = estimator or OnePageEstimator()
        self.logger = logger.bind(component="document_generator")

    # ---------------- Word ----------------

    def generate_word(
        self,
        resume: Any,
        jd: Optional[Any] = None,
        template: TemplateName = "conservative",
        *,
        strict_one_page: bool = True,
    ) -> DocumentResult:
        """生成 .docx 简历。

        Args:
            resume: dict / ResumeProfile（duck-type）
            jd: StructuredJD（用于文件名 + 可能的内容调整）
            template: "conservative" / "modern"
            strict_one_page: True → 超页抛 OnePageOverflowError

        Returns:
            DocumentResult（含 .docx bytes + 文件名 + 估算）

        Raises:
            OnePageOverflowError: strict_one_page=True 且简历超页
            DocumentGenerationError: 模板缺失或渲染失败
        """
        resume_d = self._to_dict(resume)
        jd_d = self._to_dict(jd) if jd is not None else {}

        # 1. 一页纸校验
        estimate = self.estimator.estimate(resume_d)
        if strict_one_page and estimate.overflow:
            raise OnePageOverflowError(estimate)

        # 2. 加载模板
        template_path = TEMPLATE_DIR / f"{template}.j2"
        if not template_path.exists():
            raise DocumentGenerationError(
                f"Word 模板不存在：{template_path}（可用：{sorted(p.stem for p in TEMPLATE_DIR.glob('*.j2'))}）"
            )

        # 3. 渲染 HTML（jinja2 → markdown → HTML）
        html_content = self._render_html(resume_d, jd_d, template)

        # 4. HTML → .docx（python-docx）
        try:
            from docx import Document
            from docx.shared import Pt, Mm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            raise DocumentGenerationError(
                f"python-docx 未装（pip install python-docx）：{e}"
            )

        doc = Document()
        # §1.4 边距硬约束
        for section in doc.sections:
            section.top_margin = Mm(self.A4_MARGIN_TOP_MM)
            section.bottom_margin = Mm(self.A4_MARGIN_BOTTOM_MM)
            section.left_margin = Mm(self.A4_MARGIN_LEFT_MM)
            section.right_margin = Mm(self.A4_MARGIN_RIGHT_MM)

        # 设置默认正文样式
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(self.DEFAULT_FONT_SIZE_PT)

        # 渲染 HTML → docx 段落（极简实现：把 HTML 转成 plain text + 段落切分）
        self._populate_doc_from_html(doc, html_content)

        # 5. 序列化
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        # 6. 文件名
        jd_title = jd_d.get("title", "岗位") if jd_d else "通用岗位"
        company = jd_d.get("company", "公司") if jd_d else "公司"
        filename = suggest_filename(
            name=resume_d.get("name", "简历"),
            jd_title=jd_title,
            company=company,
            ext="docx",
        )

        return DocumentResult(
            filename=filename,
            content=content,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            template_used=template,
            estimate=estimate,
        )

    # ---------------- PDF ----------------

    def generate_pdf(
        self,
        resume: Any,
        jd: Optional[Any] = None,
        template: TemplateName = "modern",
        *,
        strict_one_page: bool = True,
    ) -> DocumentResult:
        """生成 PDF 简历（HTML → playwright → bytes）。"""
        resume_d = self._to_dict(resume)
        jd_d = self._to_dict(jd) if jd is not None else {}

        # 1. 一页纸校验
        estimate = self.estimator.estimate(resume_d)
        if strict_one_page and estimate.overflow:
            raise OnePageOverflowError(estimate)

        # 2. 渲染 HTML
        html_content = self._render_html(resume_d, jd_d, template, for_pdf=True)

        # 3. HTML → PDF（playwright headless chromium via tools.generator.resume_pdf）
        pdf_bytes = html_to_pdf_safe(html_content)
        if pdf_bytes is None:
            raise DocumentGenerationError(
                "PDF 生成失败（playwright/headless chromium 不可用，"
                "或 HTML 渲染超时）。请检查 playwright install + chromium 下载状态"
            )

        # 4. 文件名
        jd_title = jd_d.get("title", "岗位") if jd_d else "通用岗位"
        company = jd_d.get("company", "公司") if jd_d else "公司"
        filename = suggest_filename(
            name=resume_d.get("name", "简历"),
            jd_title=jd_title,
            company=company,
            ext="pdf",
        )

        return DocumentResult(
            filename=filename,
            content=pdf_bytes,
            mime_type="application/pdf",
            template_used=template,
            estimate=estimate,
        )

    # ---------------- helpers ----------------

    def _render_html(
        self, resume_d: Dict[str, Any], jd_d: Dict[str, Any],
        template: TemplateName, for_pdf: bool = False,
    ) -> str:
        """jinja2 渲染模板 → HTML。"""
        try:
            from jinja2 import Environment, FileSystemLoader, select_autoescape
        except ImportError as e:
            raise DocumentGenerationError(f"jinja2 未装：{e}")

        env = Environment(
            loader=FileSystemLoader(str(TEMPLATE_DIR)),
            autoescape=select_autoescape(["html"]),
        )
        try:
            tpl = env.get_template(f"{template}.j2")
        except Exception as e:
            raise DocumentGenerationError(f"模板 {template}.j2 加载失败：{e}")

        html = tpl.render(
            resume=resume_d,
            jd=jd_d,
            styles=self._styles_for_template(template, for_pdf=for_pdf),
        )
        return html

    def _styles_for_template(self, template: TemplateName, for_pdf: bool) -> str:
        """返回模板要嵌入的 CSS。§1.4 硬约束（字号/行距/边距）。"""
        base_css = f"""
        @page {{
            size: A4;
            margin: {self.A4_MARGIN_TOP_MM}mm {self.A4_MARGIN_RIGHT_MM}mm {self.A4_MARGIN_BOTTOM_MM}mm {self.A4_MARGIN_LEFT_MM}mm;
        }}
        body {{
            font-family: "Microsoft YaHei", "PingFang SC", sans-serif;
            font-size: {self.DEFAULT_FONT_SIZE_PT}pt;
            line-height: {self.LINE_SPACING};
            color: #222;
            margin: 0;
        }}
        h1 {{ font-size: {self.SECTION_TITLE_SIZE_PT + 4}pt; margin: 8px 0 4px; }}
        h2 {{
            font-size: {self.SECTION_TITLE_SIZE_PT}pt;
            font-weight: bold;
            margin: 8px 0 4px;
            border-bottom: 1px solid #888;
        }}
        h3 {{ font-size: {self.DEFAULT_FONT_SIZE_PT + 1}pt; font-weight: bold; margin: 6px 0 2px; }}
        .contact {{ font-size: {self.DEFAULT_FONT_SIZE_PT - 1.5}pt; color: #555; }}
        .ai-template {{
            border: 2px dashed #d97706;
            padding: 4px 8px;
            background: #fffbeb;
            color: #92400e;
        }}
        ul {{ margin: 4px 0 6px 18px; padding: 0; }}
        li {{ margin-bottom: 2px; }}
        """
        if template == "modern":
            return base_css + """
            h2 { color: #2563eb; border-bottom-color: #2563eb; }
            .ai-template { border-color: #2563eb; background: #eff6ff; color: #1e40af; }
            """
        # conservative 默认灰色
        return base_css

    def _populate_doc_from_html(self, doc: Any, html: str) -> None:
        """极简 HTML → docx 段落（解析 <h1>/<h2>/<ul><li>/<div class="ai-template">）。"""
        from bs4 import BeautifulSoup
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        soup = BeautifulSoup(html, "html.parser")

        # 拿第一个 <style> 块注入（让 docx 在 Word 里打开也对齐）
        style_tag = soup.find("style")
        # 不强制注入，docx 用自己的样式即可

        # body 可能不存在（CI 最小依赖下 bs4 stub 没 body 属性）
        body = getattr(soup, "body", None) or soup
        elements = body.find_all(recursive=False) if hasattr(body, "find_all") else []
        for el in elements:
            text = el.get_text(strip=True)
            if not text:
                continue

            tag = el.name
            if tag == "h1":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(self.SECTION_TITLE_SIZE_PT + 4)
            elif tag == "h2":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(self.SECTION_TITLE_SIZE_PT)
            elif tag == "h3":
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            elif tag == "ul":
                for li in el.find_all("li", recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
            elif tag == "div" and "ai-template" in (el.get("class") or []):
                # 模式 B 虚线框标注
                p = doc.add_paragraph()
                run = p.add_run("⚠️ " + text)
                run.italic = True
                run.font.color.rgb = None  # 默认色
            else:
                # 普通段落
                doc.add_paragraph(text)

    @staticmethod
    def _to_dict(obj: Any) -> Dict[str, Any]:
        """ResumeProfile / dict / dataclass → dict（duck-type）。"""
        if obj is None:
            return {}
        if isinstance(obj, dict):
            return obj
        if hasattr(obj, "model_dump"):
            return obj.model_dump()
        if hasattr(obj, "to_db_dict"):
            return obj.to_db_dict()
        if hasattr(obj, "__dict__"):
            return vars(obj)
        return {}